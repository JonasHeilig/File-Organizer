from flask import Flask, render_template, request, redirect, url_for, send_from_directory, abort
import os
from docx import Document
from markupsafe import Markup

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/files'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB Limit


def safe_join(directory, filename):
    filepath = os.path.join(directory, filename)
    if not os.path.abspath(filepath).startswith(os.path.abspath(directory)):
        abort(403)
    return filepath


def get_document_text(file_path):
    doc = Document(file_path)
    return '<br>'.join([p.text for p in doc.paragraphs])


def save_document_text(file_path, content):
    doc = Document()
    for line in content.split('<br>'):
        doc.add_paragraph(line)
    doc.save(file_path)


def get_parent_folder(path):
    if not path:
        return ''
    return os.path.dirname(path)


@app.route('/')
@app.route('/<path:current_folder>')
def index(current_folder=''):
    base_folder = app.config['UPLOAD_FOLDER']
    current_path = os.path.join(base_folder, current_folder)

    if not os.path.exists(current_path):
        abort(404)

    parent_folder = get_parent_folder(current_folder)
    folders = [f for f in os.listdir(current_path) if os.path.isdir(os.path.join(current_path, f))]
    files = [f for f in os.listdir(current_path) if os.path.isfile(os.path.join(current_path, f))]
    return render_template('index.html', folders=folders, files=files, current_folder=current_folder,
                           parent_folder=parent_folder)


@app.route('/view/<path:current_folder>/<filename>')
def view_file(current_folder, filename):
    base_folder = app.config['UPLOAD_FOLDER']
    file_path = safe_join(os.path.join(base_folder, current_folder), filename)
    if not os.path.exists(file_path):
        abort(404)
    if filename.lower().endswith('.pdf'):
        return render_template('view_pdf.html', filename=filename)
    elif filename.lower().endswith('.doc') or filename.lower().endswith('.docx'):
        doc_content = get_document_text(file_path)
        return render_template('view_doc.html', filename=filename, doc_content=Markup(doc_content),
                               current_folder=current_folder)
    else:
        return send_from_directory(os.path.join(base_folder, current_folder), filename)


@app.route('/edit/<path:current_folder>/<filename>')
def edit_file(current_folder, filename):
    base_folder = app.config['UPLOAD_FOLDER']
    current_path = os.path.join(base_folder, current_folder)

    file_path = safe_join(current_path, filename)
    if not os.path.exists(file_path):
        abort(404)

    doc_content = get_document_text(file_path)
    return render_template('edit_doc.html', filename=filename, doc_content=Markup(doc_content),
                           current_folder=current_folder)


@app.route('/save/<path:current_folder>/<filename>', methods=['POST'])
def save_file(current_folder, filename):
    base_folder = app.config['UPLOAD_FOLDER']
    file_path = safe_join(os.path.join(base_folder, current_folder), filename)
    if not os.path.exists(file_path):
        abort(404)
    content = request.form['content']
    save_document_text(file_path, content)
    return redirect(url_for('index', current_folder=current_folder))


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    file = request.files['file']
    current_folder = request.form['current_folder']
    if file.filename == '':
        return redirect(request.url)
    if file and (file.filename.lower().endswith('.pdf') or
                 file.filename.lower().endswith('.doc') or
                 file.filename.lower().endswith('.docx')):
        filename = file.filename
        save_path = safe_join(os.path.join(app.config['UPLOAD_FOLDER'], current_folder), filename)
        file.save(save_path)
        return redirect(url_for('index', current_folder=current_folder))
    else:
        return 'Nur PDF, DOC und DOCX Dateien sind erlaubt', 400


@app.route('/create_folder', methods=['POST'])
def create_folder():
    current_folder = request.form['current_folder']
    folder_name = request.form['folder_name']
    new_folder_path = safe_join(os.path.join(app.config['UPLOAD_FOLDER'], current_folder), folder_name)
    try:
        os.makedirs(new_folder_path)
    except OSError:
        return 'Ordner konnte nicht erstellt werden', 400
    return redirect(url_for('index', current_folder=current_folder))


@app.route('/delete/<path:current_folder>/<filename>', methods=['POST'])
def delete_file(current_folder, filename):
    base_folder = app.config['UPLOAD_FOLDER']
    file_path = safe_join(os.path.join(base_folder, current_folder), filename)
    if not os.path.exists(file_path):
        abort(404)
    os.remove(file_path)
    return redirect(url_for('index', current_folder=current_folder))


@app.route('/move/<path:current_folder>/<filename>', methods=['POST'])
def move_file(current_folder, filename):
    base_folder = app.config['UPLOAD_FOLDER']
    file_path = safe_join(os.path.join(base_folder, current_folder), filename)
    if not os.path.exists(file_path):
        abort(404)
    destination_folder = request.form['destination_folder']
    destination_path = safe_join(os.path.join(base_folder, destination_folder), filename)
    os.rename(file_path, destination_path)
    return redirect(url_for('index', current_folder=current_folder))


if __name__ == '__main__':
    app.run(debug=True)
